#!/usr/bin/env python3
"""
Extract tables from DOCX documents using python-docx.
Outputs JSON format for DocSummarizer consumption.
"""

import argparse
import json
import sys
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print(json.dumps({"error": "python-docx not installed. Run: pip install python-docx"}), file=sys.stderr)
    sys.exit(1)


def extract_tables_from_docx(
    docx_path: str,
    min_rows: int = 2,
    min_cols: int = 2
) -> List[Dict[str, Any]]:
    """
    Extract tables from DOCX using python-docx.

    Args:
        docx_path: Path to DOCX file
        min_rows: Minimum rows to consider a table
        min_cols: Minimum columns to consider a table

    Returns:
        List of table dictionaries
    """
    extracted_tables = []

    try:
        doc = Document(docx_path)
    except Exception as e:
        raise ValueError(f"Failed to open DOCX: {e}")

    for table_idx, table in enumerate(doc.tables):
        # Extract table cells
        rows_data = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get cell text, handling merged cells
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            rows_data.append(row_data)

        # Filter empty rows
        non_empty_rows = [
            row for row in rows_data
            if any(cell.strip() for cell in row)
        ]

        if len(non_empty_rows) < min_rows:
            continue

        # Check column count
        col_count = len(non_empty_rows[0]) if non_empty_rows else 0
        if col_count < min_cols:
            continue

        # Detect header
        has_header = detect_header_docx(non_empty_rows, table)

        # Estimate confidence
        confidence = estimate_confidence_docx(non_empty_rows, table)

        # Get section number (approximate - based on table position)
        section = estimate_section(table_idx, len(doc.tables))

        extracted_table = {
            "page": section,  # Use section as "page" for DOCX
            "boundingBox": None,  # Not available in python-docx
            "rows": non_empty_rows,
            "hasHeader": has_header,
            "confidence": confidence,
            "metadata": {
                "tableIndex": table_idx,
                "rowCount": len(table.rows),
                "colCount": len(table.columns)
            }
        }

        extracted_tables.append(extracted_table)

    return extracted_tables


def detect_header_docx(rows: List[List[str]], table: Table) -> bool:
    """
    Detect if first row is a header in DOCX table.
    """
    if len(rows) < 2:
        return False

    # Check if first row has different formatting (python-docx limitation: style detection is complex)
    # Use simple heuristic: non-numeric cells
    first_row = rows[0]
    non_numeric_count = sum(1 for cell in first_row if not is_numeric(cell))

    return non_numeric_count >= len(first_row) * 0.6


def is_numeric(value: str) -> bool:
    """Check if string represents a number."""
    if not value or not value.strip():
        return False

    cleaned = value.replace(',', '').replace('$', '').replace('%', '').strip()

    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def estimate_confidence_docx(rows: List[List[str]], table: Table) -> float:
    """
    Estimate extraction confidence for DOCX table.
    """
    if not rows:
        return 0.0

    # Check column consistency
    col_counts = [len(row) for row in rows]
    consistent_cols = len(set(col_counts)) == 1

    # Check fill rate
    total_cells = sum(len(row) for row in rows)
    non_empty_cells = sum(1 for row in rows for cell in row if cell.strip())
    fill_rate = non_empty_cells / total_cells if total_cells > 0 else 0

    confidence = 0.6  # Base confidence (DOCX tables are more reliable than PDF)

    if consistent_cols:
        confidence += 0.3

    confidence += fill_rate * 0.1

    return min(1.0, confidence)


def estimate_section(table_idx: int, total_tables: int) -> int:
    """
    Estimate section number for table (rough approximation).
    """
    # Simple heuristic: divide document into sections
    if total_tables <= 5:
        return table_idx + 1
    else:
        return (table_idx // 5) + 1


def main():
    parser = argparse.ArgumentParser(description="Extract tables from DOCX using python-docx")
    parser.add_argument("--input", "-i", required=True, help="Input DOCX file")
    parser.add_argument("--min-rows", type=int, default=2, help="Minimum rows")
    parser.add_argument("--min-cols", type=int, default=2, help="Minimum columns")

    args = parser.parse_args()

    try:
        tables = extract_tables_from_docx(
            docx_path=args.input,
            min_rows=args.min_rows,
            min_cols=args.min_cols
        )

        # Output JSON to stdout
        print(json.dumps(tables, indent=2))

    except FileNotFoundError:
        print(json.dumps({"error": f"File not found: {args.input}"}), file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
